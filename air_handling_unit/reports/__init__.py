import os
import time
from io import BytesIO
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
import openai


class FaultCodeFiveReport:
    """Class provides the definitions for Fault Code 5 Report."""

    def __init__(
        self,
        mix_degf_err_thres: float,
        supply_degf_err_thres: float,
        delta_t_supply_fan: float,
        mat_col: str,
        sat_col: str,
        htg_vlv_col: str,
        fan_vfd_speed_col: str
    ):
        self.mix_degf_err_thres = mix_degf_err_thres
        self.supply_degf_err_thres = supply_degf_err_thres
        self.delta_t_supply_fan = delta_t_supply_fan
        self.mat_col = mat_col
        self.sat_col = sat_col
        self.htg_vlv_col = htg_vlv_col
        self.fan_vfd_speed_col = fan_vfd_speed_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc5_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 5 Plot')

        plot1a, = ax1.plot(df.index, df[self.mat_col],
                           color='g', label="Mix Temp")

        plot1b, = ax1.plot(df.index, df[self.sat_col],
                           color='b', label="Supply Temp")

        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[self.htg_vlv_col], label="Htg Valve", color="r")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('%')
        ax2.legend(loc='best')

        ax3.plot(df.index, df[output_col], label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc5_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc5_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_sat = round(
            df[self.sat_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc5_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc5_flag"

        if mat_col is None:
            mat_col = "mat"

        # calculate dataset statistics
        df["hour_of_the_day_fc5"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc5
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc5.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 5 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc5_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Five Report", 0)

        p = document.add_paragraph(
            """Fault condition five of ASHRAE Guideline 36 is (an AHU heating mode or winter time conditions only fault equation) related to flagging supply air temperatures that are out of acceptable ranges based on the mix air temperature and an assumption for heat created by the AHU supply fan in the air stream. Fault condition five equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc5_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc5_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc5_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 5 is True the average mix air temp is {flag_true_mat}°F and the outside air temp is {flag_true_sat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        # ADD in Summary Statistics
        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)
        document.add_heading('Mix Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Supply Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating the AHU temperature sensors for either the supply or mix temperature are out of calibration. Verify the mixing temperature sensor is not a probe type sensor but a long averaging type sensor that is installed properly inside the AHU mixing chamber to get a good solid true reading of the actual air mixing temperature. Poor duct design may also contribute to not having good air mixing, to troubleshoot install data loggers inside the mixing chamber or take measurements when the AHU is running of different locations in the mixing chamber to spot where better air blending needs to take place.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU temperature sensors are within calibration')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeSixReport:
    """Class provides the definitions for Fault Code 5 Report."""

    def __init__(
        self,
        vav_total_flow_col: str,
        mat_col: str,
        oat_col: str,
        rat_col: str,
        fan_vfd_speed_col: str
    ):
        self.vav_total_flow_col = vav_total_flow_col
        self.fan_vfd_speed_col = fan_vfd_speed_col
        self.mat_col = mat_col
        self.oat_col = oat_col
        self.rat_col = rat_col
        self.fan_vfd_speed_col = fan_vfd_speed_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc6_flag"

        fig, (ax1, ax2, ax3, ax4, ax5) = plt.subplots(5, 1, figsize=(25, 8))
        plt.title('Fault Conditions 6 Plot')
        ax1.plot(df.index, df['rat_minus_oat'],
                 label="Rat Minus Oat")
        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[self.vav_total_flow_col],
                 label="Total Air Flow", color="r")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('CFM')
        ax2.legend(loc='best')

        plot3a = ax3.plot(
            df.index, df['percent_oa_calc'], label="OA Frac Calc", color="m")
        plot3b = ax3.plot(df.index, df['perc_OAmin'],
                          label="OA Perc Min Calc", color="y")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('%')
        ax3.legend(loc='best')

        ax4.plot(df.index, df['percent_oa_calc_minus_perc_OAmin'],
                 label="OA Error Frac Vs Perc Min Calc", color="g")
        ax4.set_xlabel('Date')
        ax4.set_ylabel('%')
        ax4.legend(loc='best')

        ax5.plot(df.index, df[output_col], label="Fault", color="k")
        ax5.set_xlabel('Date')
        ax5.set_ylabel('Fault Flags')
        ax5.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc6_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc5_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_rat = round(
            df[self.rat_col].where(df[output_col] == 1).mean(), 2
        )

        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc5_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_rat,
            flag_true_oat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc6_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc6"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc6
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc6.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 6 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc6_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Six Report", 0)

        p = document.add_paragraph(
            """Fault condition six of ASHRAE Guideline 36 is an attempt at verifying that AHU design minimum outside air is close to the calculated outside air fraction through the outside, mix, and return air temperature sensors. A fault will get flagged in an AHU heating or mechanical cooling mode only if the calculated OA fraction is too low or too high as to compared to percent Min calculation which is the AHU total air flow divided by the design minimum outdoor air expressed as a percent. Fault condition six equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc6_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc6_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_rat,
            flag_true_oat,
            hours_motor_runtime,
            df_motor_on_filtered
        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc6_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 6 is True the average AHU mix air temperature {flag_true_mat}°F, outside air temperature is {flag_true_oat}°F, and the return air temperature is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this AHU is drawing in excessive outside air.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        # ADD in Summary Statistics
        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Mix Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Outside Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))
        
        # ADD in Summary Statistics
        document.add_heading('Return Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.rat_col].describe()))
        
        # ADD in Summary Statistics
        document.add_heading('Total Air Flow', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.vav_total_flow_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent true metric maybe yeilding sensors are out of calibration either on the AHU outside, mix, or return air temperature sensors that handle the OA fraction calculation or the totalized air flow calculation handled by a totalizing all VAV box air flows or AHU AFMS. Air flow and/or AHU temperature sensor may require recalibration.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the sensors are within calibration')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeSevenReport:
    """Class provides the definitions for Fault Code 7 Report.
        Very similar to FC 13 but uses heating valve
    """

    def __init__(
        self,
        sat_col: str,
        satsp_col: str,
        htg_col: str,
        fan_vfd_speed_col: str
    ):
        self.sat_col = sat_col
        self.satsp_col = satsp_col
        self.htg_col = htg_col
        self.fan_vfd_speed_col = fan_vfd_speed_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc7_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 7 Plot')

        plot1a, = ax1.plot(df.index, df[self.sat_col], label="SAT")
        plot1b, = ax1.plot(df.index, df[self.satsp_col], label="SATsp")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Supply Temps °F')

        ax2.plot(df.index, df[self.htg_col], color="r",
                 label="AHU Heat Vlv")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df.fc7_flag, label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc7_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc7_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_satsp = round(
            df[self.sat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_sat = round(
            df[self.satsp_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc7_mode,
            percent_true,
            percent_false,
            flag_true_satsp,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc7_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc7"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc7
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc7.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 7 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc7_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Seven Report", 0)

        p = document.add_paragraph(
            """Fault condition seven of ASHRAE Guideline 36 is an AHU heating mode only with an attempt at verifying an AHU heating or cooling valve is not stuck or leaking by verifying AHU supply temperature to supply temperature setpoint. Fault condition seven equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc7_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc7_mode,
            percent_true,
            percent_false,
            flag_true_satsp,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc7_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 7 is True the average AHU supply air setpoint is {flag_true_satsp} in °F and the supply air temperature is {flag_true_sat} in °F.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        # ADD in Summary Statistics
        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp Setpoint', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.satsp_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Heating Coil Valve', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.htg_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating the AHU heating valve maybe broken or there could be a flow issue with the amount of hot water flowing through the coil or that the boiler system reset is too aggressive and there isnt enough heat being produced by this coil. It could be worth viewing mechanical blue prints for this AHU design schedule to see what hot water temperature this coil was designed for and compare it to actual hot water supply temperatures. IE., an AHU hot water coil sized to have a 180°F water flowing through it may have a durastic reduction in performance the colder the hot water is flowing through it, if need be consult a mechanical design engineer to rectify.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU heating valve operates Ok.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeEightReport:
    """Class provides the definitions for Fault Code 8 Report."""

    def __init__(
        self,
        sat_col: str,
        mat_col: str,
        fan_vfd_speed_col: str,
        economizer_sig_col: str,
    ):
        self.sat_col = sat_col
        self.mat_col = mat_col
        self.fan_vfd_speed_col = fan_vfd_speed_col
        self.economizer_sig_col = economizer_sig_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc8_flag"

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 8 Plot')

        plot1a, = ax1.plot(df.index, df[self.sat_col], label="SAT")
        plot1b, = ax1.plot(df.index, df[self.mat_col], label="MAT")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Temps °F')

        ax2.plot(df.index, df.fc8_flag, label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc8_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc8_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_sat = round(
            df[self.sat_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc8_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc8_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc8"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc8
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc8.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 8 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc8_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Eight Report", 0)

        p = document.add_paragraph(
            """Fault condition Eight of ASHRAE Guideline 36 is an AHU economizer free cooling mode only with an attempt at flagging conditions when the AHU mixing air temperature the supply air temperature are not approximately equal. Fault condition eight equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc8_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc8_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc8_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 8 is True the average AHU mix air is {flag_true_mat} in °F and the supply air temperature is {flag_true_sat} in °F.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Mix Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating temperature sensor error or the heating/cooling coils are leaking potentially creating simultenious heating/cooling which can be an energy penalty for running the AHU in this fashion. Verify AHU mix/supply temperature sensor calibration in addition to a potential mechanical issue of a leaking valve. A leaking valve can be troubleshot by isolating the valve closed by manual shut off valves where piping lines enter the AHU coil and then verifying any changes in the AHU discharge air temperature.')
        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU components are within calibration for this fault equation Ok.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeNineReport:
    """Class provides the definitions for Fault Code 9 Report."""

    def __init__(
        self,
        satsp_col: str,
        oat_col: str,
        fan_vfd_speed_col: str,
        economizer_sig_col: str,
    ):

        self.satsp_col = satsp_col
        self.oat_col = oat_col
        self.fan_vfd_speed_col = fan_vfd_speed_col
        self.economizer_sig_col = economizer_sig_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc9_flag"

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 10 Plot')

        plot1a, = ax1.plot(df.index, df[self.satsp_col], label="SATSP")
        plot1b, = ax1.plot(df.index, df[self.oat_col], label="OAT")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Temps °F')

        ax2.plot(df.index, df.fc9_flag, label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc9_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc9_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_satsp = round(
            df[self.satsp_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc9_mode,
            percent_true,
            percent_false,
            flag_true_oat,
            flag_true_satsp,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc9_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc9"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc10
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc9.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 9 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc9_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Nine Report", 0)

        p = document.add_paragraph(
            """Fault condition nine of ASHRAE Guideline 36 is an AHU economizer free cooling mode only with an attempt at flagging conditions where the outside air temperature is too warm for cooling without additional mechanical cooling. Fault condition nine equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc9_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc9_mode,
            percent_true,
            percent_false,
            flag_true_oat,
            flag_true_satsp,
            hours_motor_runtime,
            df_motor_on_filtered

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc9_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 9 is True the average outside air is {flag_true_oat} in °F and the supply air temperature setpoinht is {flag_true_satsp} in °F.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        # ADD in Summary Statistics
        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp Setpoint', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.satsp_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Outside Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating temperature sensor error or the cooling valve is stuck open or leaking causing overcooling. Trouble shoot a leaking valve by isolating the coil with manual shutoff valves and verify a change in AHU discharge air temperature with the AHU running.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU components are within calibration for this fault equation Ok.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeTenReport:
    """Class provides the definitions for Fault Code 10 Report."""

    def __init__(
        self,
        oat_col: str,
        mat_col: str,
        clg_col: str,
        economizer_sig_col: str,
        fan_vfd_speed_col: str,
    ):
        self.oat_col = oat_col
        self.mat_col = mat_col
        self.clg_col = clg_col
        self.economizer_sig_col = economizer_sig_col
        self.fan_vfd_speed_col = fan_vfd_speed_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc10_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 10 Plot')

        plot1a, = ax1.plot(df.index, df[self.mat_col], label="MAT")
        plot1b, = ax1.plot(df.index, df[self.oat_col], label="OAT")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Temps °F')

        plot2a, = ax2.plot(
            df.index, df[self.clg_col], label="AHU Cool Vlv", color="r")
        plot2b, = ax2.plot(
            df.index, df[self.economizer_sig_col], label="AHU Dpr Cmd", color="g")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df.fc10_flag, label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc10_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc10_mode = (delta * df[output_col]
                           ).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc10_mode,
            percent_true,
            percent_false,
            flag_true_oat,
            flag_true_mat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc10_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc10"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc10
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc10.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 10 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc10_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Ten Report", 0)

        p = document.add_paragraph(
            """Fault condition ten of ASHRAE Guideline 36 is an AHU economizer + mechanical cooling mode only with an attempt at flagging conditions where the outside air temperature and mixing air temperatures are not approximetely equal when the AHU is in a 100% outside air mode. Fault condition ten equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc10_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc10_mode,
            percent_true,
            percent_false,
            flag_true_oat,
            flag_true_mat,
            hours_motor_runtime,
            df_motor_on_filtered
        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc10_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 9 is True the average outside air is {flag_true_oat} in °F and the mixing air temperature is {flag_true_mat} in °F.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        # ADD in Summary Statistics
        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Mixing Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Outside Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating temperature sensor error or the mixing air dampers are stuck or broken with the inability for the AHU to go into a proper 100 percent outside air mode. If the outside air temperature is a global variable on the BAS verify (IE, installed to the boiler plant controller and then shared via supervisory level logic on the BAS to the AHU controllers on the BAS network) that where the actual OA temperature is installed that is on the North side of the building in the shade. On the AHU verify mix temperature sensor calibration and that the mixing dampers have good proper rotation with good seals when in the closed position. When testing AHU systems operating in a 100 percent outside air mode it could be worth verifying exhaust systems or return fans are operating properly. In thoery if alot of air is being pumped into the building and it is allowed to be exhaust or relieved properly, a balanced building will not have any issues of closing or opening egress doors to the building due to excess positive building pressure.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU components are within calibration for this fault equation Ok.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeElevenReport:
    """Class provides the definitions for Fault Code 11 Report."""

    def __init__(
        self,
        sat_sp_col: str,
        oat_col: str,
        clg_col: str,
        economizer_sig_col: str,
        fan_vfd_speed_col: str,
    ):
        self.sat_sp_col = sat_sp_col
        self.oat_col = oat_col
        self.clg_col = clg_col
        self.economizer_sig_col = economizer_sig_col
        self.fan_vfd_speed_col = fan_vfd_speed_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc11_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 11 Plot')

        plot1a, = ax1.plot(df.index, df[self.sat_sp_col], label="SATSP")
        plot1b, = ax1.plot(df.index, df[self.oat_col], label="OAT")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Temps °F')

        plot2a, = ax2.plot(
            df.index, df[self.clg_col], label="AHU Cool Vlv", color="r")
        plot2b, = ax2.plot(
            df.index, df[self.economizer_sig_col], label="AHU Dpr Cmd", color="g")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df.fc11_flag, label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc11_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc11_mode = (delta * df[output_col]
                           ).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_sat_sp = round(
            df[self.sat_sp_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc11_mode,
            percent_true,
            percent_false,
            flag_true_oat,
            flag_true_sat_sp,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc11_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc11"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc11
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc11.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 11 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc11_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Eleven Report", 0)

        p = document.add_paragraph(
            """Fault condition eleven of ASHRAE Guideline 36 is an AHU economizer + mechanical cooling mode only with an attempt at flagging conditions where the outside air temperature is too low for 100% outside air AHU operating mode. Fault condition Eleven equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc11_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc11_mode,
            percent_true,
            percent_false,
            flag_true_oat,
            flag_true_sat_sp,
            hours_motor_runtime,
            df_motor_on_filtered

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc11_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 11 is True the average AHU mix air is {flag_true_oat} in °F and the supply air temperature is {flag_true_sat_sp} in °F.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        # ADD in Summary Statistics
        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp Setpoint', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            str(df_motor_on_filtered[self.sat_sp_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Outside Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating temperature sensor error or the heating coil could be leaking potentially creating simultenious heating/cooling scenorio which can be an energy penalty for running the AHU in this fashion. Also visually verify with the AHU off via lock-out-tag-out that the mixing dampers operates effectively. To do this have one person the BAS sending operator override commands to drive the damper back and forth. The other person should put on eyes on the operation of the actuator motor driving the OA dampers 100 percent open and then closed and visually verify the dampers rotate effectively per BAS command where to also visually verify the dampers have a good seal when in the closed position. Also consider looking into BAS programming that may need tuning or parameter adjustments for the staging between OS state changes between AHU modes of operation.')
        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU components are within calibration for this fault equation Ok.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeTwelveReport:
    """Class provides the definitions for Fault Code 12 Report."""

    def __init__(
        self,
        sat_col: str,
        mat_col: str,
        clg_col: str,
        economizer_sig_col: str,
        fan_vfd_speed_col: str,
    ):
        self.sat_col = sat_col
        self.mat_col = mat_col
        self.clg_col = clg_col
        self.economizer_sig_col = economizer_sig_col
        self.fan_vfd_speed_col = fan_vfd_speed_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc12_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 12 Plot')

        plot1a, = ax1.plot(df.index, df[self.sat_col], label="SAT")
        plot1b, = ax1.plot(df.index, df[self.mat_col], label="MAT")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Temps °F')

        plot2a, = ax2.plot(
            df.index, df[self.clg_col], label="AHU Cool Vlv", color="r")
        plot2b, = ax2.plot(
            df.index, df[self.economizer_sig_col], label="AHU Dpr Cmd", color="g")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df.fc12_flag, label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc12_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc12_mode = (delta * df[output_col]
                           ).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_sat = round(
            df[self.sat_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc12_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc12_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc12"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc12
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc12.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 12 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc12_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Twelve Report", 0)

        p = document.add_paragraph(
            """Fault condition Twelve of ASHRAE Guideline 36 is an AHU economizer + mechanical cooling mode and AHU mechanical cooling mode only with an attempt at flagging conditions when the AHU mixing air temperature is warmer than the supply air temperature. Fault condition Twelve equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc12_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc12_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc12_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 12 is True the average AHU mix air is {flag_true_mat} in °F and the supply air temperature is {flag_true_sat} in °F.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        # ADD in Summary Statistics
        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Mix Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating temperature sensor error or the heating/cooling coils are leaking potentially creating simultenious heating/cooling which can be an energy penalty for running the AHU in this fashion. Verify AHU mix/supply temperature sensor calibration in addition to a potential mechanical issue of a leaking valve. A leaking valve can be troubleshot by isolating the valve closed by manual shut off valves where piping lines enter the AHU coil and then verifying any changes in the AHU discharge air temperature.')
        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU components are within calibration for this fault equation Ok.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeThirteenReport:
    """Class provides the definitions for Fault Code 13 Report.
        Verify similar to FC 7 but uses cooling valve
    """

    def __init__(
        self,
        sat_col: str,
        satsp_col: str,
        clg_col: str,
        economizer_sig_col: str,
        fan_vfd_speed_col: str
    ):
        self.sat_col = sat_col
        self.satsp_col = satsp_col
        self.clg_col = clg_col
        self.economizer_sig_col = economizer_sig_col
        self.fan_vfd_speed_col = fan_vfd_speed_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc13_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 13 Plot')

        plot1a, = ax1.plot(df.index, df[self.sat_col], label="SAT")
        plot1b, = ax1.plot(df.index, df[self.satsp_col], label="SATsp")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Supply Temps °F')

        plot2a, = ax2.plot(
            df.index, df[self.clg_col], label="AHU Cool Vlv", color="r")
        plot2b, = ax2.plot(
            df.index, df[self.economizer_sig_col], label="AHU Dpr Cmd", color="g")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df.fc13_flag, label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc13_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc13_mode = (delta * df[output_col]
                           ).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_satsp = round(
            df[self.satsp_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_sat = round(
            df[self.sat_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.fan_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc13_mode,
            percent_true,
            percent_false,
            flag_true_satsp,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        vav_total_flow: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc13_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc13"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc13
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc13.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 13 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc13_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Thirteen Report", 0)

        p = document.add_paragraph(
            """Fault condition thirteen of ASHRAE Guideline 36 is an AHU cooling mode only with an attempt at verifying an AHU cooling valve is not stuck or leaking by verifying AHU supply temperature to supply temperature setpoint. Fault condition thirteen equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc13_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc13_mode,
            percent_true,
            percent_false,
            flag_true_satsp,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc13_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 13 is True the average AHU supply air setpoint is {flag_true_satsp} in °F and the supply air temperature is {flag_true_sat} in °F.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Supply Air Temp Setpoint', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.satsp_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Cooling Coil Valve', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.clg_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating the AHU cooling valve maybe broken or there could be a flow issue with the amount of cold water flowing through the coil or that the chiller system leaving temperature reset is too aggressive and there isnt enough cold air being produced by this cooling coil. If this AHU has a DX cooling coil there could be a problem with the refrigerant charge. It could be worth viewing mechanical blue prints for this AHU design schedule to see what cold water temperature this coil was designed for and compare it to actual cold water supply temperatures. IE., an AHU cooling coil sized to have a 44°F water flowing through it may have significant performance reduction with 48°F water flowing through it and under design day type high load conditions this AHU may not meet setpoint or properly dehumidify the air for the building which could potentially also lead to IAQ or mold issues if %RH levels in the zones are kept within tollerance. Also check excessive outside air faults in fault condition 6 that the AHU isnt taking in too much outdoor air which could also cause coil performance issues if the load on the coil is higher than what it was intended for.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU cooling valve operates Ok.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document
